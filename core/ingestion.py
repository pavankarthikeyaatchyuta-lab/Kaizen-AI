"""
Kaizen AI — Document Ingestion Pipeline v1.1
Turning Industrial Documents into Operational Intelligence.

Handles: PDF, DOCX, Excel, Images (scanned)
Enhancements: Hierarchical chunking, version detection,
              OCR confidence propagation, extensible metadata,
              Document Trust Score
"""

import os
import re
import json
import hashlib
import logging
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

import PyPDF2
import docx
import openpyxl

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ─── Data Structures ──────────────────────────────────────────────────────────

@dataclass
class DocumentChunk:
    chunk_id: str
    doc_id: str
    source_file: str
    doc_type: str
    content: str
    metadata: dict = field(default_factory=dict)
    # metadata keys (all optional, added as discovered):
    # page, chapter, section, subsection, heading,
    # ocr, ocr_confidence, language, tables_detected,
    # figures_detected, equipment_refs, revision, doc_date


@dataclass
class ExtractedEntity:
    entity_type: str
    value: str
    context: str
    source_chunk_id: str
    confidence: float


@dataclass
class DocumentMeta:
    """Document-level metadata — version, trust, fingerprint."""
    doc_id: str
    filename: str
    doc_type: str
    checksum: str
    revision: Optional[str]        # "Rev3", "v2", "2025-01"
    equipment_id: Optional[str]    # "P-104", "M-201"
    doc_date: Optional[str]
    trust_score: float             # 0–1
    trust_factors: dict            # breakdown for explainability
    is_duplicate: bool = False
    superseded_by: Optional[str] = None


@dataclass
class IngestionResult:
    doc_id: str
    filename: str
    doc_type: str
    doc_meta: Optional[DocumentMeta]
    chunks: list
    entities: list
    success: bool
    error: Optional[str] = None


# ─── Document Type Classifier ─────────────────────────────────────────────────

DOC_TYPE_KEYWORDS = {
    "manual":           ["oem", "manual", "operating instruction", "user guide", "specification"],
    "maintenance_log":  ["maintenance", "work order", "repair", "service", "lubrication", "overhaul"],
    "inspection":       ["inspection", "audit", "checklist", "survey", "condition report"],
    "incident":         ["incident", "accident", "near miss", "failure report", "breakdown", "root cause"],
    "sop":              ["procedure", "sop", "standard operating", "safety", "permit to work"],
    "compliance":       ["factory act", "oisd", "peso", "iso", "dgms", "compliance", "regulation"],
}

def classify_document(text: str, filename: str) -> str:
    text_lower = (text[:2000] + " " + filename).lower()
    scores = {dtype: sum(1 for kw in kws if kw in text_lower)
              for dtype, kws in DOC_TYPE_KEYWORDS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "general"


# ─── Version & Equipment Detection ───────────────────────────────────────────

REVISION_PATTERNS = [
    r'[Rr]ev(?:ision)?\.?\s*(\d+|[A-Z])',
    r'[Vv](?:ersion)?\.?\s*(\d+(?:\.\d+)?)',
    r'v(\d+\.\d+)',
    r'Ed(?:ition)?\.?\s*(\d+)',
    r'(\d{4}-\d{2}-\d{2})',   # date as version
]

EQUIPMENT_PATTERNS = [
    r'\b([A-Z]{1,3}-\d{2,4}[A-Z]?)\b',   # P-104, M-201A
    r'\b([A-Z]{1,3}\d{3,5})\b',           # P104, M2014
]

def detect_revision(text: str, filename: str) -> Optional[str]:
    combined = filename + " " + text[:1000]
    for pat in REVISION_PATTERNS:
        m = re.search(pat, combined)
        if m:
            return m.group(0).strip()
    return None

def detect_equipment_id(text: str, filename: str) -> Optional[str]:
    combined = filename + " " + text[:500]
    for pat in EQUIPMENT_PATTERNS:
        m = re.search(pat, combined)
        if m:
            return m.group(1)
    return None

def compute_checksum(filepath: str) -> str:
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


# ─── Document Trust Score ─────────────────────────────────────────────────────

def compute_trust_score(
    ocr_used: bool,
    avg_ocr_confidence: float,
    revision_detected: bool,
    parse_success: bool,
    chunk_count: int,
    is_duplicate: bool,
) -> tuple[float, dict]:
    """
    Trust score = weighted combination of quality signals.
    Returned alongside a breakdown dict for explainable AI output.
    """
    factors = {
        "parse_success":       1.0 if parse_success else 0.0,
        "ocr_quality":         avg_ocr_confidence if ocr_used else 1.0,
        "revision_known":      1.0 if revision_detected else 0.85,
        "content_richness":    min(1.0, chunk_count / 10),
        "not_duplicate":       0.6 if is_duplicate else 1.0,
    }
    weights = {
        "parse_success":    0.30,
        "ocr_quality":      0.30,
        "revision_known":   0.15,
        "content_richness": 0.15,
        "not_duplicate":    0.10,
    }
    score = sum(factors[k] * weights[k] for k in factors)
    return round(score, 3), factors


# ─── Entity Extraction ────────────────────────────────────────────────────────

ENTITY_PATTERNS = {
    "equipment": [
        r'\b([A-Z]{1,4}[-\s]?\d{2,4}[A-Z]?)\b',
    ],
    "component": [
        r'\b(bearing|seal|impeller|shaft|coupling|rotor|gasket|'
        r'flange|valve|actuator|sensor|piston|diaphragm)\b',
    ],
    "standard": [
        r'\b(OISD[-\s]?\d+|Factory Act|PESO|ISO[-\s]?\d+|'
        r'DGMS|IS[-\s]?\d+|API[-\s]?\d+)\b',
    ],
    "failure_type": [
        r'\b(vibration|corrosion|erosion|cavitation|overheating|'
        r'leakage|crack|fracture|wear|fatigue|misalignment|imbalance|'
        r'contamination|blockage|seizure)\b',
    ],
    "measurement": [
        r'\b(\d+(?:\.\d+)?)\s*(mm/s|rpm|°C|bar|psi|kPa|Hz|kW|A|V|mm|m³/h)\b',
    ],
    "maintenance_action": [
        r'\b(replace|lubricate|inspect|overhaul|calibrate|align|'
        r'clean|tighten|adjust|flush|test|verify)\b',
    ],
}

def extract_entities(text: str, chunk_id: str, base_confidence: float = 0.85) -> list:
    entities = []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    for sentence in sentences:
        for etype, patterns in ENTITY_PATTERNS.items():
            for pattern in patterns:
                for match in re.findall(pattern, sentence, re.IGNORECASE):
                    value = match if isinstance(match, str) else " ".join(match).strip()
                    if len(value) > 1:
                        entities.append(ExtractedEntity(
                            entity_type=etype,
                            value=value.strip(),
                            context=sentence.strip()[:200],
                            source_chunk_id=chunk_id,
                            confidence=round(base_confidence, 3),
                        ))
    return entities


# ─── Section Header Detector ──────────────────────────────────────────────────

SECTION_HEADER_RE = re.compile(
    r'^(\d{1,2}(?:\.\d{1,2}){0,3})\s+([A-Z][A-Za-z\s\-&/]{3,80})$',
    re.MULTILINE
)

def parse_section_headers(text: str) -> list[tuple[str, str, int]]:
    """Returns list of (section_number, heading_text, char_offset)."""
    results = []
    for m in SECTION_HEADER_RE.finditer(text):
        results.append((m.group(1), m.group(2).strip(), m.start()))
    return results

def section_depth(number: str) -> tuple:
    """'3.2.1' → ('3', '3.2', '3.2.1')"""
    parts = number.split(".")
    chapter   = parts[0] if len(parts) >= 1 else ""
    section   = ".".join(parts[:2]) if len(parts) >= 2 else ""
    subsection = ".".join(parts[:3]) if len(parts) >= 3 else ""
    return chapter, section, subsection


# ─── PDF Extractor ────────────────────────────────────────────────────────────

def extract_pdf(filepath: str, doc_id: str) -> tuple[list, dict]:
    import PyPDF2
    chunks = []
    filename = Path(filepath).name
    full_text_parts = []
    try:
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                raw_text = page.extract_text()
                if raw_text and raw_text.strip():
                    full_text_parts.append(raw_text.strip())
            
        full_text = "\n".join(full_text_parts)
        doc_type = classify_document(full_text[:2000], filename)
        headers = parse_section_headers(full_text)

        if headers:
            for i, (sec_num, heading, offset) in enumerate(headers):
                next_offset = headers[i+1][2] if i+1 < len(headers) else len(full_text)
                content = full_text[offset:next_offset].strip()
                if len(content) < 30:
                    continue
                chapter, section, subsection = section_depth(sec_num)
                chunk_id = f"{doc_id}_s{sec_num.replace('.','_')}"
                approx_page = max(1, int((offset / max(len(full_text), 1)) * num_pages) + 1)
                
                chunks.append(DocumentChunk(
                    chunk_id=chunk_id, doc_id=doc_id, source_file=filename,
                    doc_type=doc_type, content=content[:2000],
                    metadata={
                        "page": approx_page, "chapter": chapter, "section": section,
                        "subsection": subsection, "heading": heading, "section_number": sec_num,
                        "ocr": False, "ocr_confidence": 1.0, "chunk_confidence": 1.0,
                    }
                ))
        else:
            word_list = full_text.split()
            chunk_size = 400
            for i in range(0, len(word_list), chunk_size):
                content = " ".join(word_list[i:i+chunk_size])
                approx_page = max(1, int((i / max(len(word_list), 1)) * num_pages) + 1)
                chunk_id = f"{doc_id}_p{approx_page}_c{i//chunk_size}"
                chunks.append(DocumentChunk(
                    chunk_id=chunk_id, doc_id=doc_id, source_file=filename,
                    doc_type=doc_type, content=content,
                    metadata={"page": approx_page, "ocr": False, "ocr_confidence": 1.0, "chunk_confidence": 1.0}
                ))
        return chunks, {"ocr_used": False, "avg_ocr_confidence": 1.0, "page_count": num_pages}
    except Exception as e:
        logger.error(f"PDF extraction failed for {filename}: {e}")
        return [], {"ocr_used": False, "avg_ocr_confidence": 0.0, "page_count": 0}


# ─── DOCX Extractor ───────────────────────────────────────────────────────────

def extract_docx(filepath: str, doc_id: str) -> tuple[list, dict]:
    filename = Path(filepath).name
    chunks = []
    try:
        document = docx.Document(filepath)
        full_text = "\n".join(p.text for p in document.paragraphs)
        doc_type = classify_document(full_text[:2000], filename)

        current_heading = "Introduction"
        current_sec_num = ""
        current_content = []

        for para in document.paragraphs:
            if para.style.name.startswith("Heading"):
                if current_content:
                    content = " ".join(current_content)
                    chunk_id = f"{doc_id}_{current_sec_num or current_heading[:15].replace(' ','_')}"
                    _, section, subsection = section_depth(current_sec_num) if current_sec_num else ("", "", "")
                    chunks.append(DocumentChunk(
                        chunk_id=chunk_id, doc_id=doc_id, source_file=filename,
                        doc_type=doc_type, content=content,
                        metadata={
                            "heading": current_heading,
                            "section": section,
                            "subsection": subsection,
                            "ocr": False,
                            "ocr_confidence": 1.0,
                            "chunk_confidence": 1.0,
                        }
                    ))
                # Detect section number from heading
                m = re.match(r'^(\d[\d.]*)\s+(.*)', para.text.strip())
                current_sec_num = m.group(1) if m else ""
                current_heading = m.group(2).strip() if m else para.text.strip()
                current_content = []
            elif para.text.strip():
                current_content.append(para.text.strip())

        if current_content:
            chunk_id = f"{doc_id}_{current_heading[:15].replace(' ','_')}_end"
            chunks.append(DocumentChunk(
                chunk_id=chunk_id, doc_id=doc_id, source_file=filename,
                doc_type=doc_type, content=" ".join(current_content),
                metadata={"heading": current_heading, "ocr": False,
                          "ocr_confidence": 1.0, "chunk_confidence": 1.0}
            ))

        return chunks, {"ocr_used": False, "avg_ocr_confidence": 1.0, "page_count": 0}

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return [], {"ocr_used": False, "avg_ocr_confidence": 0.0, "page_count": 0}


# ─── Excel Extractor ──────────────────────────────────────────────────────────

def extract_excel(filepath: str, doc_id: str) -> tuple[list, dict]:
    import openpyxl
    filename = Path(filepath).name
    chunks = []
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content = ""
            for row in sheet.iter_rows(values_only=True):
                content += " | ".join(str(v) for v in row if v is not None) + "\n"
            chunk_id = f"{doc_id}_{sheet_name[:12]}"
            chunks.append(DocumentChunk(
                chunk_id=chunk_id, doc_id=doc_id, source_file=filename,
                doc_type=classify_document(sheet_name + content[:300], filename),
                content=content[:2000],
                metadata={"section": sheet_name, "ocr": False, "ocr_confidence": 1.0, "chunk_confidence": 1.0}
            ))
        return chunks, {"ocr_used": False, "avg_ocr_confidence": 1.0, "page_count": 0}
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return [], {"ocr_used": False, "avg_ocr_confidence": 0.0, "page_count": 0}


# ─── Image Extractor ─────────────────────────────────────────────────────────

def extract_image(filepath: str, doc_id: str) -> tuple[list, dict]:
    filename = Path(filepath).name
    logger.warning("Image extraction not supported in serverless mode.")
    return [], {"ocr_used": False, "avg_ocr_confidence": 0.0, "page_count": 0}


# ─── Version Deduplication Registry ──────────────────────────────────────────

class DocumentRegistry:
    """
    In-memory registry for version/duplicate detection.
    In production: persist to DB.
    """
    def __init__(self):
        # equipment_id -> list of (doc_id, revision, filename)
        self._index: dict[str, list] = {}

    def register(self, doc_meta: DocumentMeta) -> DocumentMeta:
        eid = doc_meta.equipment_id or "unknown"
        existing = self._index.get(eid, [])

        # Check for exact duplicate (same checksum)
        for prev in existing:
            if prev["checksum"] == doc_meta.checksum:
                doc_meta.is_duplicate = True
                logger.warning(f"Duplicate detected: {doc_meta.filename} matches {prev['filename']}")
                return doc_meta

        # Check for older revision of same equipment
        if existing and doc_meta.revision:
            for prev in existing:
                if prev.get("revision") and prev["revision"] != doc_meta.revision:
                    # Mark the older one as superseded
                    logger.info(
                        f"Version update: {prev['filename']} superseded by {doc_meta.filename}"
                    )
                    # In production: update prev doc in DB
                    prev["superseded_by"] = doc_meta.doc_id

        existing.append({
            "doc_id": doc_meta.doc_id,
            "filename": doc_meta.filename,
            "revision": doc_meta.revision,
            "checksum": doc_meta.checksum,
            "superseded_by": None,
        })
        self._index[eid] = existing
        return doc_meta

    def get_latest(self, equipment_id: str) -> Optional[dict]:
        docs = self._index.get(equipment_id, [])
        active = [d for d in docs if not d.get("superseded_by")]
        return active[-1] if active else None

    def get_superseded_warning(self, equipment_id: str) -> Optional[str]:
        docs = self._index.get(equipment_id, [])
        superseded = [d for d in docs if d.get("superseded_by")]
        if superseded:
            return (f"⚠ {len(superseded)} superseded version(s) for equipment {equipment_id}. "
                    f"Using latest revision for reasoning.")
        return None


# ─── Main Ingestion Entry Point ───────────────────────────────────────────────

EXTRACTORS = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_excel,
    ".xls":  extract_excel,
    ".png":  extract_image,
    ".jpg":  extract_image,
    ".jpeg": extract_image,
    ".tiff": extract_image,
}

# Global registry (pass your own instance for multi-session use)
_default_registry = DocumentRegistry()

def ingest_document(filepath: str, registry: DocumentRegistry = None) -> IngestionResult:
    """
    Main entry point. Pass any supported industrial document.
    Returns structured chunks + extracted entities + document metadata.
    """
    if registry is None:
        registry = _default_registry

    filepath = str(filepath)
    filename = Path(filepath).name
    ext = Path(filepath).suffix.lower()
    doc_id = f"doc_{abs(hash(filepath + filename)) % 100000:05d}"

    logger.info(f"Ingesting: {filename}")

    extractor = EXTRACTORS.get(ext)
    if not extractor:
        return IngestionResult(
            doc_id=doc_id, filename=filename, doc_type="unknown",
            doc_meta=None, chunks=[], entities=[], success=False,
            error=f"Unsupported file type: {ext}"
        )

    chunks, stats = extractor(filepath, doc_id)

    if not chunks:
        return IngestionResult(
            doc_id=doc_id, filename=filename, doc_type="unknown",
            doc_meta=None, chunks=[], entities=[], success=False,
            error="Extraction returned no content"
        )

    # Combine text sample for metadata detection
    sample_text = " ".join(c.content[:300] for c in chunks[:3])
    doc_type = chunks[0].doc_type

    # Build document metadata
    revision    = detect_revision(sample_text, filename)
    equipment_id = detect_equipment_id(sample_text, filename)
    checksum    = compute_checksum(filepath)
    trust_score, trust_factors = compute_trust_score(
        ocr_used=stats["ocr_used"],
        avg_ocr_confidence=stats["avg_ocr_confidence"],
        revision_detected=revision is not None,
        parse_success=True,
        chunk_count=len(chunks),
        is_duplicate=False,
    )

    doc_meta = DocumentMeta(
        doc_id=doc_id, filename=filename, doc_type=doc_type,
        checksum=checksum, revision=revision, equipment_id=equipment_id,
        doc_date=None, trust_score=trust_score, trust_factors=trust_factors,
    )

    # Register for version/duplicate detection
    doc_meta = registry.register(doc_meta)

    # Inject trust + revision into each chunk's metadata
    for chunk in chunks:
        chunk.metadata.update({
            "doc_trust_score": trust_score,
            "revision": revision,
            "equipment_id": equipment_id,
        })

    # Extract entities (confidence scaled by chunk's OCR confidence)
    all_entities = []
    for chunk in chunks:
        base_conf = chunk.metadata.get("chunk_confidence", 0.85)
        entities = extract_entities(chunk.content, chunk.chunk_id, base_confidence=base_conf)
        all_entities.extend(entities)

    logger.info(
        f"✓ {filename} | type={doc_type} | rev={revision} | "
        f"chunks={len(chunks)} | entities={len(all_entities)} | "
        f"trust={trust_score:.2f} | ocr={stats['ocr_used']}"
    )

    return IngestionResult(
        doc_id=doc_id, filename=filename, doc_type=doc_type,
        doc_meta=doc_meta, chunks=chunks, entities=all_entities, success=True
    )


def ingest_folder(folder_path: str, registry: DocumentRegistry = None) -> list:
    """Batch ingest all supported documents in a folder."""
    if registry is None:
        registry = _default_registry

    folder = Path(folder_path)
    files = [f for f in folder.iterdir() if f.suffix.lower() in EXTRACTORS]
    logger.info(f"Found {len(files)} documents in {folder_path}")

    results = [ingest_document(str(f), registry) for f in files]
    success = sum(1 for r in results if r.success)
    logger.info(f"Ingestion complete: {success}/{len(results)} successful")
    return results


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = ingest_document(sys.argv[1])
        print(json.dumps({
            "doc_id": result.doc_id,
            "doc_type": result.doc_type,
            "revision": result.doc_meta.revision if result.doc_meta else None,
            "equipment_id": result.doc_meta.equipment_id if result.doc_meta else None,
            "trust_score": result.doc_meta.trust_score if result.doc_meta else None,
            "is_duplicate": result.doc_meta.is_duplicate if result.doc_meta else None,
            "chunks": len(result.chunks),
            "entities": len(result.entities),
            "sample_chunk": {
                "content": result.chunks[0].content[:300],
                "metadata": result.chunks[0].metadata,
            } if result.chunks else None,
            "sample_entities": [
                {"type": e.entity_type, "value": e.value, "confidence": e.confidence}
                for e in result.entities[:10]
            ],
        }, indent=2))
    else:
        print("Usage: python ingestion.py <path_to_document>")
